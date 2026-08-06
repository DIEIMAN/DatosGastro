#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Edición editable (.docx) de la edición de conducción del Atlas DGDGAS.

Qué produce
-----------
    outputs/polos_gastro/ATLAS_V2/ATLAS_REFERENCIAS_GASTRONOMICAS_CABA_DGDGAS.docx
    outputs/polos_gastro/ATLAS_V2/ATLAS_REFERENCIAS_GASTRONOMICAS_CABA_DGDGAS_DOCX.pdf

El .pdf de al lado es la **prueba de paginación** del .docx: se dibuja desde el mismo
modelo de documento, con la misma tipografía (Calibri del sistema), la misma caja de
texto (A4, márgenes de 2 cm) y las mismas alturas de imagen. No lo produce Word —esta
máquina no tiene Word ni LibreOffice—, pero la paginación no depende del motor de
maquetado: cada página del .docx está delimitada por un salto de página explícito y su
contenido se mide antes de escribirlo, con holgura reservada.

Las dos ediciones existentes (conducción y técnica, en PDF) NO se tocan: este generador
sólo lee del corpus cerrado y de la cartografía vectorial, y escribe dos archivos nuevos.

Regla de las imágenes
---------------------
Cada mapa se vuelve a dibujar **desde el renderer vectorial** (`cartografia_vectorial_v2`)
dentro de una caja del tamaño que necesita el documento, y recién ahí se rasteriza a PNG
a 200 dpi. No se recorta ninguna página del PDF: recortar es lo que partía los mapas.

`Vista` escala la geometría para que entre completa en la caja y la centra; ensanchar la
caja no achica el dibujo, agrega contexto alrededor. Por eso el encuadre se elige así:

    k        = escala máxima que entra en el área disponible de la página
    w_geo    = ancho que ocupa la geometría a esa escala   (idem alto)
    caja     = w_geo/h_geo estirado hasta entrar en la banda de proporciones cómodas
               [R_MIN, R_MAX], sin pasar del área disponible

Resultado: la escala del mapa es siempre la máxima posible para la página, y las
proporciones extremas (R02 y el contexto Corrientes-Abasto, horizontales; R04, R14 y R15,
verticales) se compensan con margen de contexto, nunca recortando contenido.

Reproducción (sin red, desde la raíz del repositorio):

    .venv/Scripts/python.exe -B outputs/polos_gastro/ATLAS_V2/scripts/build_atlas_docx.py
"""

from __future__ import annotations

import argparse
import hashlib
import json
import math
import os
import re
import shutil
import sys
import zipfile
from dataclasses import dataclass, field
from pathlib import Path

SCRIPT = Path(__file__).resolve()
SCRIPTS = SCRIPT.parent
if str(SCRIPTS) not in sys.path:
    sys.path.insert(0, str(SCRIPTS))

import fitz                                   # PyMuPDF
from PIL import Image
from reportlab.lib.colors import HexColor
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.pdfgen import canvas as rl_canvas

import docx
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import (WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING,
                            WD_TAB_ALIGNMENT, WD_TAB_LEADER)
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

import build_atlas_v2 as B
import cartografia_vectorial_v2 as CV
import contenido_conduccion as C
import lenguaje_conduccion as LENG

# =======================================================================================
# Geometría de la página y escala tipográfica
# =======================================================================================

CM = 28.3464567                     # puntos por centímetro
PAGE_W_CM, PAGE_H_CM = 21.0, 29.7
MARGIN_CM = 2.0
COL_W_CM = PAGE_W_CM - 2 * MARGIN_CM            # 17,0 cm
COL_H_CM = PAGE_H_CM - 2 * MARGIN_CM            # 25,7 cm
COL_W = COL_W_CM * CM
COL_H = COL_H_CM * CM

# Holgura reservada al pie de cada página: absorbe cualquier diferencia entre el corte de
# línea de esta medición y el de Word. Un renglón de cuerpo mide ~0,45 cm.
HOLGURA_CM = 1.2

# Interlineado simple de Calibri: (ascender + descender + lineGap) / unitsPerEm en `hhea`.
CALIBRI_SINGLE = 1.2207

AZUL = "1F3B57"          # institucional, títulos
AZUL_CLARO = "2C7FB8"    # títulos de tercer nivel
GRIS_TEXTO = "222222"    # cuerpo
GRIS = "555555"          # textos secundarios
NARANJA = "C0762B"       # advertencias, y sólo advertencias
GRIS_LINEA = "D9DEE5"    # filetes de la tabla

FAMILIA_COLOR = dict(CV.FAMILIA_COLOR)

# Estilos. `word` es el nombre del estilo de Word que se aplica al párrafo; los estilos
# propios se crean en `preparar_estilos`. `mult` es el interlineado múltiple de Word.
STYLES: dict[str, dict] = {
    "Portada.marca":   dict(word="Portada marca",     size=15, bold=True,  color=AZUL,       before=0,  after=2,  mult=1.0),
    "Portada.org":     dict(word="Portada organismo", size=11, bold=False, color=GRIS,       before=0,  after=26, mult=1.0),
    "Portada.titulo":  dict(word="Title",             size=26, bold=True,  color=AZUL,       before=0,  after=8,  mult=1.0),
    "Portada.sub":     dict(word="Subtitle",          size=14, bold=False, color=GRIS,       before=0,  after=22, mult=1.05),
    "Portada.cuerpo":  dict(word="Normal",            size=11.5, bold=False, color=GRIS,     before=0,  after=14, mult=1.15),
    "Portada.rotulo":  dict(word="Portada rotulo",    size=11, bold=True,  color=AZUL,       before=8,  after=3,  mult=1.0),
    "Portada.nota":    dict(word="Advertencia",       size=10, bold=False, color=NARANJA,    before=14, after=4,  mult=1.15),

    "H1":              dict(word="Heading 1",         size=18, bold=True,  color=AZUL,       before=0,  after=9,  mult=1.0),
    "H2":              dict(word="Heading 2",         size=13, bold=True,  color=AZUL,       before=11, after=4,  mult=1.0),
    "H3":              dict(word="Heading 3",         size=11.5, bold=True, color=AZUL_CLARO, before=9, after=3,  mult=1.0),

    "Normal":          dict(word="Normal",            size=10.5, bold=False, color=GRIS_TEXTO, before=0, after=6, mult=1.15),
    "Bloque":          dict(word="Bloque de ficha",   size=10.5, bold=True,  color=AZUL,       before=8, after=2, mult=1.0),
    "Sec":             dict(word="Texto secundario",  size=9.5, bold=False, color=GRIS,       before=0,  after=5,  mult=1.15),
    "SecIt":           dict(word="Texto secundario cursiva", size=9.5, bold=False, italic=True, color=GRIS, before=0, after=7, mult=1.15),
    "Warn":            dict(word="Advertencia",       size=9.5, bold=False, color=NARANJA,    before=6,  after=5,  mult=1.15),
    "Pie":             dict(word="Pie de mapa",       size=9.5, bold=False, color=GRIS,       before=0,  after=3,  mult=1.12),
    "Bullet":          dict(word="List Bullet",       size=10.5, bold=False, color=GRIS_TEXTO, before=0, after=3, mult=1.12),
    "BulletSec":       dict(word="List Bullet Sec",   size=9.5, bold=False, color=GRIS,       before=0,  after=3,  mult=1.12),
    "Imagen":          dict(word="Imagen de mapa",    size=10.5, bold=False, color=GRIS_TEXTO, before=5, after=6, mult=1.0),
    "Tabla":           dict(word="Celda de tabla",    size=9.5, bold=False, color=GRIS_TEXTO, before=0, after=0, mult=1.08),
    "TablaCab":        dict(word="Cabecera de tabla", size=9.5, bold=True,  color=AZUL,       before=0,  after=0,  mult=1.08),

    "Indice":          dict(word="Indice seccion",    size=11, bold=True,  color=AZUL,       before=7,  after=2,  mult=1.0),
    "IndiceZona":      dict(word="Indice zona",       size=10, bold=False, color=GRIS_TEXTO, before=0,  after=1,  mult=1.0),
    "PiePagina":       dict(word="Footer",            size=9,  bold=False, color=GRIS,       before=0,  after=0,  mult=1.0),
}

SANGRIA_BULLET_CM = 0.63

# =======================================================================================
# Modelo de documento
# =======================================================================================


@dataclass
class Run:
    text: str
    bold: bool = False
    italic: bool = False
    color: str | None = None

    def __post_init__(self) -> None:
        self.text = llano(self.text)


@dataclass
class P:
    runs: list[Run]
    style: str = "Normal"
    keep_with_next: bool = False


@dataclass
class Bul:
    items: list[list[Run]]
    style: str = "Bullet"


@dataclass
class Img:
    path: Path
    w_cm: float
    h_cm: float


@dataclass
class Tbl:
    headers: list[str]
    rows: list[list[str]]
    widths_cm: list[float]

    def __post_init__(self) -> None:
        self.headers = [llano(h) for h in self.headers]
        self.rows = [[llano(celda) for celda in fila] for fila in self.rows]


@dataclass
class Salto:
    """Salto de página explícito."""


@dataclass
class Idx:
    """Una línea del índice: título, puntos suspensivos y número de página."""
    texto: str
    pagina: int
    style: str = "Indice"


@dataclass
class MarcaIndice:
    """Dónde va el índice. `armar_indice` lo reemplaza por sus páginas reales."""


# Ancho que se le reserva al número de página al final de cada línea del índice.
ANCHO_FOLIO = 1.0 * CM


def txt(s: str, style: str = "Normal", **kw) -> P:
    return P([Run(s)], style, **kw)


# =======================================================================================
# Métrica tipográfica (Calibri real, la misma que usa Word)
# =======================================================================================

FONTS = {
    "r":  "Calibri",
    "b":  "Calibri-Bold",
    "i":  "Calibri-Italic",
    "bi": "Calibri-BoldItalic",
}


def registrar_calibri() -> None:
    base = Path(os.environ.get("SystemRoot", r"C:\Windows")) / "Fonts"
    archivos = {"r": "calibri.ttf", "b": "calibrib.ttf", "i": "calibrii.ttf", "bi": "calibriz.ttf"}
    for clave, nombre in archivos.items():
        ruta = base / nombre
        if not ruta.is_file():
            raise FileNotFoundError(
                f"Falta {nombre}: la prueba de paginación necesita Calibri del sistema.")
        pdfmetrics.registerFont(TTFont(FONTS[clave], str(ruta)))


def fuente(bold: bool, italic: bool) -> str:
    return FONTS["bi" if (bold and italic) else "b" if bold else "i" if italic else "r"]


def ancho(texto: str, bold: bool, italic: bool, size: float) -> float:
    return pdfmetrics.stringWidth(texto, fuente(bold, italic), size)


def partir(runs: list[Run], width: float, st: dict) -> list[list[tuple[str, bool, bool, str]]]:
    """Corta una secuencia de runs en líneas que entran en `width`.

    Devuelve, por línea, los segmentos (texto, negrita, cursiva, color) que la componen.
    El criterio de corte es el de Word: se corta entre palabras, y una palabra sola más
    ancha que la columna se deja salir (no se parte).
    """
    size = st["size"]
    lineas: list[list[tuple[str, bool, bool, str]]] = [[]]
    usado = 0.0
    for run in runs:
        color = run.color or st["color"]
        bold = run.bold or st.get("bold", False)
        italic = run.italic or st.get("italic", False)
        piezas = run.text.split(" ")
        for idx, palabra in enumerate(piezas):
            if idx:
                palabra_con_espacio = " " + palabra
            else:
                palabra_con_espacio = palabra
            if not palabra_con_espacio:
                continue
            w = ancho(palabra_con_espacio, bold, italic, size)
            if usado > 0 and usado + w > width:
                sin_espacio = palabra_con_espacio.lstrip(" ")
                lineas.append([])
                usado = 0.0
                palabra_con_espacio = sin_espacio
                w = ancho(sin_espacio, bold, italic, size)
            if not palabra_con_espacio:
                continue
            if lineas[-1] and lineas[-1][-1][1] == bold and lineas[-1][-1][2] == italic \
                    and lineas[-1][-1][3] == color:
                anterior = lineas[-1][-1]
                lineas[-1][-1] = (anterior[0] + palabra_con_espacio, bold, italic, color)
            else:
                lineas[-1].append((palabra_con_espacio, bold, italic, color))
            usado += w
    return [ln for ln in lineas if ln] or [[("", False, False, st["color"])]]


def alto_lineas(n: int, st: dict) -> float:
    return n * st["size"] * CALIBRI_SINGLE * st["mult"]


def alto_bloque(bloque, width: float = COL_W) -> float:
    """Alto en puntos que ocupa un bloque, espacios antes/después incluidos."""
    if isinstance(bloque, (Salto, MarcaIndice)):
        return 0.0
    if isinstance(bloque, Idx):
        st = STYLES[bloque.style]
        lineas = partir([Run(bloque.texto)], width - ANCHO_FOLIO, st)
        return alto_lineas(len(lineas), st) + st["before"] + st["after"]
    if isinstance(bloque, Img):
        st = STYLES["Imagen"]
        return bloque.h_cm * CM + st["before"] + st["after"]
    if isinstance(bloque, P):
        st = STYLES[bloque.style]
        return alto_lineas(len(partir(bloque.runs, width, st)), st) + st["before"] + st["after"]
    if isinstance(bloque, Bul):
        # Los estilos de lista de Word llevan `contextualSpacing`: entre dos viñetas
        # consecutivas del mismo estilo NO se aplica el espacio posterior. Sólo la última
        # de la lista lo lleva. Medirlo de otra manera sobreestima el alto de la página.
        st = STYLES[bloque.style]
        w = width - SANGRIA_BULLET_CM * CM
        total = sum(alto_lineas(len(partir(item, w, st)), st) for item in bloque.items)
        return total + st["before"] + st["after"]
    if isinstance(bloque, Tbl):
        return alto_tabla(bloque)
    raise TypeError(bloque)


PAD_CELDA = 3.2          # relleno vertical de celda, por lado


def alto_fila(valores: list[str], anchos: list[float], estilo: str) -> float:
    st = STYLES[estilo]
    lineas = max(len(partir([Run(v)], w - 2 * 4.0, st)) for v, w in zip(valores, anchos))
    return alto_lineas(lineas, st) + 2 * PAD_CELDA


def alto_tabla(t: Tbl) -> float:
    anchos = [w * CM for w in t.widths_cm]
    total = alto_fila(t.headers, anchos, "TablaCab") + 2
    for fila in t.rows:
        total += alto_fila(fila, anchos, "Tabla")
    return total


# =======================================================================================
# Cartografía: re-render vectorial a PNG
# =======================================================================================

# Banda de proporciones (ancho/alto) admitidas para la caja de la imagen. Fuera de ella,
# la caja se estira —agregando contexto, nunca recortando— hasta volver a entrar.
R_MIN, R_MAX = 0.75, 1.18

_ULTIMA_VISTA: dict[str, CV.Vista] = {}
_VISTA_INIT = CV.Vista.__init__


def _vista_init(self, bounds, x, y, w, h):
    _VISTA_INIT(self, bounds, x, y, w, h)
    _ULTIMA_VISTA["v"] = self


CV.Vista.__init__ = _vista_init


# El control de bordes del renderer mide el ancho del rótulo en horizontal aunque el
# rótulo se dibuje rotado. Un nombre de avenida casi vertical ocupa a lo alto lo que mide
# a lo ancho: en la caja del PDF (18 x 22,7 cm) entraba igual, y en la del documento
# —más baja— el marco lo cortaba. «Av. Intendente Bullrich» en el mapa de Las Cañitas
# salía partido por eso.
#
# La corrección vive acá, en el generador del documento, y no en `cartografia_vectorial_v2`:
# ese módulo produce las dos ediciones en PDF ya publicadas y no se toca. El rótulo se
# corre hacia adentro lo justo para entrar; no se suprime ninguno, porque cada mapa tiene
# que rotular lo que su pie promete.
_MARCO: dict[str, tuple[float, float, float, float]] = {}
_HALO_ORIGINAL = CV.texto_con_halo


def _halo_acotado(c, x, y, texto, fuente_mapa, tamano, color, halo=None, grosor=1.5,
                  centrado=True, angulo=0.0):
    from reportlab.lib.colors import white as _blanco
    marco = _MARCO.get("rect")
    if marco and centrado:
        ancho_txt = pdfmetrics.stringWidth(texto, fuente_mapa, tamano)
        rad = math.radians(angulo)
        semi_x = abs(ancho_txt / 2 * math.cos(rad)) + abs(tamano / 2 * math.sin(rad))
        semi_y = abs(ancho_txt / 2 * math.sin(rad)) + abs(tamano / 2 * math.cos(rad))
        bx, by, bw, bh = marco
        if 2 * semi_x + 6 <= bw:
            x = min(max(x, bx + semi_x + 3), bx + bw - semi_x - 3)
        if 2 * semi_y + 6 <= bh:
            y = min(max(y, by + semi_y + 3), by + bh - semi_y - 3)
    return _HALO_ORIGINAL(c, x, y, texto, fuente_mapa, tamano, color,
                          _blanco if halo is None else halo, grosor, centrado, angulo)


CV.texto_con_halo = _halo_acotado


def dibujante(clave: str, tipo: str):
    """Devuelve la función que dibuja ese mapa sobre un canvas, en la caja dada."""
    if tipo == "principal":
        return lambda c, x, y, w, h: B.CARTOGRAFO.mapa_referencia(c, clave, x, y, w, h)
    if tipo == "ampliacion":
        return lambda c, x, y, w, h: B.CARTOGRAFO.mapa_detalle(
            c, clave, x, y, w, h, B.EXTRAS_DETALLE)
    if tipo == "general":
        return lambda c, x, y, w, h: B.CARTOGRAFO.mapa_general(c, x, y, w, h, B.REF_PAGES)
    raise ValueError(tipo)


def proporcion_geografica(dibuja, tmp: Path) -> tuple[float, float]:
    """Ancho y alto, en metros, del encuadre que el propio renderer elige."""
    c = rl_canvas.Canvas(str(tmp), pagesize=(400, 400))
    dibuja(c, 0.0, 0.0, 400.0, 400.0)
    c.showPage()
    c.save()
    minx, miny, maxx, maxy = _ULTIMA_VISTA["v"].bounds_geo
    return maxx - minx, maxy - miny


def encuadre(dx: float, dy: float, w_max_cm: float, h_max_cm: float) -> tuple[float, float]:
    """Caja de la imagen, en cm, para ese encuadre geográfico.

    Primero la escala máxima que entra en el área disponible; después la caja se estira
    hacia una proporción cómoda. Estirar la caja NO achica el mapa: `Vista` mantiene la
    escala y rellena con contexto.
    """
    k = min(w_max_cm / dx, h_max_cm / dy)
    w_geo, h_geo = dx * k, dy * k
    w_box = min(w_max_cm, max(w_geo, R_MIN * h_geo))
    h_box = min(h_max_cm, max(h_geo, w_box / R_MAX))
    w_box = min(w_max_cm, max(w_box, R_MIN * h_box))
    return round(w_box, 3), round(h_box, 3)


def render_mapa(clave: str, tipo: str, w_cm: float, h_cm: float, dpi: int,
                destino: Path, tmp: Path) -> dict:
    """Dibuja el mapa vectorial en una caja de w_cm x h_cm y lo rasteriza a PNG."""
    dibuja = dibujante(clave, tipo)
    w_pt, h_pt = w_cm * CM, h_cm * CM
    c = rl_canvas.Canvas(str(tmp), pagesize=(w_pt, h_pt))
    # 1 pt de retiro: el marco del mapa se dibuja sobre el borde de la caja y, pegado al
    # borde de la página, la mitad del trazo quedaría fuera del PNG.
    _MARCO["rect"] = (1.0, 1.0, w_pt - 2.0, h_pt - 2.0)
    try:
        info = dibuja(c, 1.0, 1.0, w_pt - 2.0, h_pt - 2.0)
    finally:
        _MARCO.pop("rect", None)
    c.showPage()
    c.save()
    doc = fitz.open(tmp)
    pix = doc[0].get_pixmap(dpi=dpi, alpha=False)
    img = Image.frombytes("RGB", (pix.width, pix.height), pix.samples)
    doc.close()
    img.save(destino, format="PNG", optimize=True)
    return {"archivo": destino.name, "clave": clave, "tipo": tipo,
            "ancho_cm": round(w_cm, 2), "alto_cm": round(h_cm, 2),
            "px": f"{pix.width}x{pix.height}", "dpi": dpi,
            "escala_m": info.get("escala_m", ""),
            "bytes": destino.stat().st_size}


# =======================================================================================
# Contenido: se lee del corpus de la edición de conducción, sin reescribir
# =======================================================================================

SECCION_ZONAS_TITULO = "Las 22 zonas, una por una"

COMO_SE_CONSTRUYERON_TITULO = "Cómo se construyeron las zonas"

# Texto entregado por la Dirección para esta edición.
#
# PENDIENTE DE CONFIRMACIÓN DE LA DIRECCIÓN (2026-08-05). Hasta el 2026-08-04 este bloque se
# transcribía tal cual. La corrección de esa noche reescribió la intro y dos pasajes más
# —Patricia Pecora lo comentó con «luego vemos la redacción y explicarlo de otra forma», o
# sea que ella misma lo dejó abierto—. Se aplica, pero el texto de origen ya no es idéntico
# al entregado: la Dirección tiene que confirmar la nueva redacción antes de darlo por
# cerrado, y el comentario de Patricia sigue abierto en el Doc.
COMO_SE_CONSTRUYERON: list[tuple[str, object]] = [
    ("p", "La lista de las 22 zonas la definió la Dirección a partir del territorio que ya "
          "venía siguiendo; no la produjo un cálculo automático sobre toda la Ciudad. Sobre "
          "esa lista se trabajó zona por zona, en tres pasos:"),
    ("h", "Primero, ubicar la oferta."),
    ("p", "Se relevaron los locales gastronómicos de cada zona con la fuente disponible para "
          "ese lugar (el relevamiento propio de la Dirección o un directorio comercial en "
          "línea) y se contó una sola vez a los que aparecían repetidos. "
          "No todas las zonas se pudieron relevar de la misma manera, y por eso el Atlas "
          "distingue cuatro situaciones:"),
    ("b", ["Zonas con relevamiento propio de la Dirección.",
           "Zonas relevadas sobre un directorio comercial en línea.",
           "Zonas donde el relevamiento llegó a su tope y el número quedó como un mínimo: "
           "se escriben con «al menos».",
           "Zonas que se pudieron describir y ubicar, pero sin un conteo propio."]),
    ("p", "Esa diferencia es la razón por la que los números de una zona no se comparan "
          "con los de otra ni se suman entre sí. Un número más alto no significa más "
          "oferta: significa que se contó de otra manera."),
    ("h", "Segundo, dibujar el área."),
    ("p", "El contorno de cada zona se construyó a partir de dónde están los locales "
          "relevados, no a partir de un límite administrativo. Por eso una zona puede no "
          "coincidir con el barrio que lleva su nombre. Donde dos zonas vecinas quedaban "
          "superpuestas, el área se repartió entre ambas para que ninguna quede contenida "
          "dentro de la otra, respetando las separaciones que la Dirección definió."),
    ("p", "El tamaño del área no indica cantidad de oferta: depende de cuán repartidos "
          "estén los locales. Una zona con pocos locales muy dispersos puede dibujarse más "
          "grande que una con muchos locales juntos."),
    ("h", "Tercero, escribir la lectura."),
    ("p", "Cada zona tiene una ficha que explica qué es ese lugar, cuánta oferta se "
          "relevó, de qué partes se compone y qué hay que tener en cuenta al leerla. Las "
          "partes internas —como Palermo Soho o el Barrio Chino— son componentes de su "
          "zona, no zonas nuevas."),
    ("h", "Qué no dicen estos números"),
    ("p", "Los números cuentan locales relevados, no locales abiertos hoy. No son un "
          "padrón, no miden facturación, empleo ni superficie, y no ordenan las zonas de "
          "mejor a peor. Los contornos muestran dónde se encontró la oferta; ninguno fue "
          "adoptado como límite oficial."),
]

NOTA_EDITABLE = ("Documento editable: los ajustes pueden hacerse directamente sobre este "
                 "archivo.")
# Gemela de `render_conduccion`: misma corrección editorial del 2026-08-04, porque las dos
# son la línea de conducción y no pueden decir cosas distintas.
CIERRE_COMO_LEER = ("Este Atlas dice dónde hay oferta gastronómica en la Ciudad y cómo es "
                    "cada lugar. Cada zona se lee por separado, con lo que explica su propia "
                    "ficha.")
SUBTITULO_MAPA_GENERAL = "Cada forma es una zona; el color indica su familia."

# El subtítulo del corpus —«22 lecturas territoriales y cartográficas»— no le dice nada a
# quien abre el documento por primera vez. Acá se dice lo mismo con las palabras que el
# propio Atlas usa en la página siguiente.
SUBTITULO_PORTADA = "22 zonas gastronómicas de la Ciudad, descriptas y mapeadas una por una"

# Los códigos R01 a R22 encabezan cada ficha, cada mapa y cada fila de la tabla final, y
# no se explicaban en ninguna parte. Se aclaran donde aparecen por primera vez.
SUBTITULO_FAMILIAS = ("Cinco formas de estar en el territorio. El color de cada familia es "
                      "el que usa el mapa de la página siguiente. Cada zona lleva además "
                      "un código, de R01 a R22, para poder nombrarla y encontrarla: es una "
                      "referencia, no un orden de importancia.")

# Las cinco definiciones estaban escritas desde adentro del trabajo. Se dicen con lo que
# el Atlas ya explica en las fichas de cada familia.
DEFINICION_FAMILIA = {
    "polo": "Una concentración de oferta reconocible, con identidad propia.",
    "multiparte": "Un polo formado por partes separadas, que el Atlas muestra separadas "
                  "porque así están en el territorio.",
    "eje": "La oferta se organiza a lo largo de una calle o avenida, no en una mancha.",
    "segmentada": "No funciona como un solo lugar: son varias áreas vecinas que conviene "
                  "mirar por separado.",
    "dispersa": "Hay oferta comprobada, pero repartida: no se encontró un centro que la "
                "organice.",
}


# =======================================================================================
# Lectura sin contexto · reescrituras de superficie
#
# Una lectura completa del .docx, hecha desde el lugar de quien lo abre por primera vez y
# no trabajó en esto, dejó una lista de frases que obligaban a volver atrás, de palabras
# que nadie usa en una conversación y de frases que admitían una lectura equivocada. Acá
# se corrige CÓMO se dice cada una. QUÉ se dice no cambia: ninguna cifra se mueve, ninguna
# salvedad desaparece y ninguna geometría se toca.
#
# El corpus cerrado y los dos módulos que producen las ediciones ya publicadas no se
# editan: esto es una capa local del .docx, con la frase original como clave. Si el corpus
# cambia y una clave deja de existir, el control `reescrituras_aplicadas` falla en vez de
# dejar una reescritura muerta y silenciosa; `diff_de_cifras_contra_conduccion` compara
# par por par los números de cada frase antes y después.
#
# El registro legible de todo esto, con el motivo de cada cambio, es
# `qa/LECTURA_SIN_CONTEXTO.md`.
# =======================================================================================

REESCRITURAS: dict[str, str] = {

    # -- Portada y resumen --------------------------------------------------------------

    # La portada atribuía todo a un relevamiento propio, y tres páginas después el propio
    # Atlas dice que en varias zonas se usó un directorio comercial en línea.
    C.FUENTE_UNICA:
        "Fuente: relevamiento de la Dirección General de Desarrollo Gastronómico, con la "
        "información disponible para cada zona. Los datos corresponden a julio de 2026.",

    # `C.NOTA_GRUPOS` ya no se reescribe acá. La corrección editorial del 2026-08-04 la
    # resolvió en el origen —explica las cuatro formas de relevo, que era la observación de
    # Patricia— y esa redacción sirve igual para el PDF y para el .docx. Mantener además una
    # reescritura propia era justamente lo que hacía que las dos ediciones de conducción
    # dijeran cosas distintas.

    # Listas de nombre + número sin unidad: la primera de cada línea la dice.
    "Caballito 907 · Villa Crespo 646 · Chacarita 327 · Boulevard Caseros 66":
        "Caballito 907 locales · Villa Crespo 646 · Chacarita 327 · Boulevard Caseros 66",
    "Villa Urquiza 189 · Devoto 119 · Avenida Boedo 79 · Donado–Holmberg 40":
        "Villa Urquiza 189 locales · Devoto 119 · Avenida Boedo 79 · Donado–Holmberg 40",
    "Centro y Microcentro 797 · Abasto 314 · La Paternal 254 · Esmeralda–Paraguay 216 · "
    "Federico Lacroze 211 · Villa Pueyrredón 158 · García del Río 40":
        "Centro y Microcentro 797 locales · Abasto 314 · La Paternal 254 · "
        "Esmeralda–Paraguay 216 · Federico Lacroze 211 · Villa Pueyrredón 158 · "
        "García del Río 40",

    # El grupo listaba cinco zonas y en la nota aparecían dos más, con sus dos números
    # pegados al final sin decir cuál era de cuál.
    "Están descriptas y ubicadas. Belgrano y Costanera Norte solo tienen datos de un "
    "relevamiento anterior: 697 y 72.":
        "Están descriptas y ubicadas, sin un número propio. Belgrano y Costanera Norte "
        "tampoco tienen conteo propio: solo hay datos de un relevamiento anterior "
        "(Belgrano, 697 locales; Costanera Norte, 72).",

    # `C.CAJA_USO` ya no se reescribe acá. Esta entrada existía para sacar la referencia a
    # «la página 6», que en un documento editable no se puede seguir; la corrección del
    # 2026-08-04 la sacó en el origen y además explicó las áreas en positivo, así que el
    # texto de conducción ya sirve tal cual para el .docx.

    # -- La frase de cifra de las cinco zonas sin conteo --------------------------------

    # «Zona caracterizada» se lee como una excusa y «caracterizada» no es una palabra de
    # conversación. Lo que se hizo va primero, y lo que falta se nombra sin rodeos.
    "Zona caracterizada, sin conteo propio": "Descripta y ubicada, sin conteo de locales",

    # -- Tabla final --------------------------------------------------------------------

    "Ref.": "Código",

    # -- Pie de los 31 mapas ------------------------------------------------------------

    "Trazo punteado: es el área con que se consultó la zona, no el contorno de la oferta "
    "encontrada.":
        "Trazo punteado: marca el área dentro de la cual se buscó, no la forma de la "
        "oferta encontrada.",
    "Números grises sobre el fondo: las 15 comunas. Trazo punteado: es el área con que se "
    "consultó la zona, no el contorno de la oferta encontrada.":
        "Números grises sobre el fondo: las 15 comunas. Trazo punteado: marca el área "
        "dentro de la cual se buscó, no la forma de la oferta encontrada.",

    # -- «La lectura» como sujeto: es la zona la que no llega hasta ahí -----------------

    "La lectura no se extiende a todo el sur de la Ciudad.":
        "Esta zona no abarca todo el sur de la Ciudad.",
    "La lectura no se extiende más allá del frente relevado ni se une con el Centro.":
        "La zona no sigue más allá del frente relevado y no se une con el Centro.",
    "La lectura no se extiende a Parque Patricios ni al resto de Barracas.":
        "La zona no llega a Parque Patricios ni al resto de Barracas.",
    "La delimitación es provisoria y la lectura no abarca todo Balvanera ni todo Almagro.":
        "La delimitación es provisoria y la zona no abarca todo Balvanera ni todo Almagro.",
    "La "
    "evidencia no alcanza para una avenida gastronómica continua, y la lectura no se "
    "extiende a Parque Patricios ni a Barracas.":
        ""
        "La evidencia no alcanza para una avenida gastronómica continua, y la zona no "
        "llega a Parque Patricios ni a Barracas.",
    "La "
    "periferia del polo no está definida con precisión y la lectura no se extiende a "
    "Villa Pueyrredón ni a Villa del Parque.":
        ""
        "La periferia del polo no está definida con precisión y la zona no llega a Villa "
        "Pueyrredón ni a Villa del Parque.",

    # -- R04 · Puerto Madero -------------------------------------------------------------

    "Puerto Madero ordena su oferta gastronómica a lo largo de los diques, sobre dos "
    "frentes enfrentados y con hitos repartidos en el borde. Es una banda larga antes "
    "que una mancha alrededor de un centro.":
        "Puerto Madero ordena su oferta gastronómica a lo largo de los diques, sobre las "
        "dos orillas enfrentadas, con hitos repartidos en el borde. Es una banda larga y "
        "no una mancha alrededor de un centro.",
    "El "
    "área excluye el espejo de agua de los cuatro diques, descontado con cartografía "
    "pública de acceso abierto verificada contra la capa oficial de cuerpos de agua de "
    "la Ciudad.":
        ""
        "El área deja afuera el agua de los cuatro diques, descontada con cartografía "
        "pública verificada contra el mapa oficial de cuerpos de agua de la Ciudad.",

    # -- R05 · Belgrano ------------------------------------------------------------------

    "Belgrano es un polo de escala grande, con tres centros gastronómicos distintos: el "
    "Barrio Chino, el Bajo Belgrano y Belgrano R. Los tres funcionan dentro de una misma "
    "zona y con pesos distintos.":
        "Belgrano es un polo de escala grande, con tres centros gastronómicos distintos: "
        "el Barrio Chino, el Bajo Belgrano y Belgrano R. Los tres funcionan dentro de una "
        "misma zona, y no todos con la misma importancia.",
    # La ficha nombraba los tres centros de una manera y el mapa de la página siguiente,
    # de otra. Los dos nombres son del corpus: se dicen juntos una vez.
    "Barrio Chino, Bajo Belgrano y Belgrano R.":
        "Barrio Chino (en el mapa, Barrio Chino–Belgrano C), Bajo Belgrano "
        "(Cabildo–Juramento) y Belgrano R.",
    "Los "
    "697 no equivalen a locales abiertos hoy ni se ponen al lado de las cifras de otras "
    "zonas.":
        ""
        "Los 697 no equivalen a locales abiertos hoy, y no sirven para decir si Belgrano "
        "tiene más oferta que otra zona.",

    # -- R06 · Recoleta ------------------------------------------------------------------

    "Un relevamiento anterior dejó datos de trabajo sobre Recoleta, pero no un número "
    "publicable para la zona.":
        "Un relevamiento anterior dejó datos sobre Recoleta, pero no un número que pueda "
        "publicarse para toda la zona.",

    # -- R09 y R10 · lo dibujado no es todo lo contado -----------------------------------

    "Se contó una sola vez a los locales que aparecían repetidos. El mapa dibuja los dos "
    "focos, que no cubren la totalidad de lo relevado.":
        "Se contó una sola vez a los locales que aparecían repetidos. El mapa dibuja los "
        "dos focos; parte de los locales relevados queda fuera de ellos.",
    "327 locales relevados en Chacarita, contando una sola vez los repetidos. El mapa "
    "dibuja los dos focos, que no cubren todo lo relevado. No equivale a locales "
    "abiertos hoy.":
        "327 locales relevados en Chacarita, contando una sola vez los repetidos. El mapa "
        "dibuja los dos focos; parte de esos locales queda fuera de ellos. No equivale a "
        "locales abiertos hoy.",
    "Se contó una sola vez a los locales que aparecían repetidos. El mapa dibuja los dos "
    "núcleos, que no cubren la totalidad de lo relevado.":
        "Se contó una sola vez a los locales que aparecían repetidos. El mapa dibuja los "
        "dos núcleos; parte de los locales relevados queda fuera de ellos.",
    "907 locales relevados en Caballito, contando una sola vez los repetidos. El mapa "
    "dibuja los dos núcleos, que no cubren todo lo relevado. No equivale a locales "
    "abiertos hoy.":
        "907 locales relevados en Caballito, contando una sola vez los repetidos. El mapa "
        "dibuja los dos núcleos; parte de esos locales queda fuera de ellos. No equivale "
        "a locales abiertos hoy.",
    "El Patio de los Lecheros se miró como punto de control y quedó fuera de la zona.":
        "El Patio de los Lecheros se revisó y quedó fuera de la zona.",
    "Parque Rivadavia salió de la lectura vigente y no se vuelve a usar.":
        "Parque Rivadavia se descartó y ya no forma parte de esta zona.",
    "Los "
    "núcleos son independientes y no forman una sola área; Parque Rivadavia quedó fuera "
    "de la lectura vigente.":
        ""
        "Los núcleos son independientes y no forman una sola área; Parque Rivadavia quedó "
        "fuera de la zona.",

    # -- R11 · Boulevard Caseros ---------------------------------------------------------

    "47 de los locales quedaron a menos de 250 metros del eje: es una descripción, no la "
    "prueba de una avenida continua.":
        "47 de esos locales están a menos de 250 metros del boulevard: es una "
        "descripción, no la prueba de que la oferta sea continua.",

    # -- R12 · Centro/Microcentro --------------------------------------------------------

    "Tribunales se lee aparte y con menos respaldo que las demás.":
        "Tribunales se mira aparte y con menos información que las demás.",

    # -- R17 · Villa Urquiza -------------------------------------------------------------

    "El mapa ubica el de Triunvirato, el único con recorrido cerrado en el material "
    "disponible; Monroe y Congreso aparecen como calles de referencia.":
        "El mapa dibuja el de Triunvirato, el único del que se tiene el recorrido "
        "completo; Monroe y Congreso aparecen sólo como calles de referencia.",
    "El área aproximada de Villa Urquiza. De sus tres calles —Triunvirato, Monroe y "
    "Congreso— el mapa ubica la de Triunvirato, la única con recorrido cerrado en el "
    "material disponible.":
        "El área aproximada de Villa Urquiza. De sus tres calles —Triunvirato, Monroe y "
        "Congreso— el mapa dibuja la de Triunvirato, la única de la que se tiene el "
        "recorrido completo.",

    # -- R18 · Esmeralda–Paraguay --------------------------------------------------------

    "La mayor cantidad de locales no está en el centro del círculo sino en las bandas "
    "intermedias, entre 100 y 300 metros.":
        "La mayor cantidad de locales no está junto al cruce, sino entre 100 y 300 metros "
        "de distancia.",
    # Un radio no es un centro: la frase negaba una cosa distinta de la que nombraba.
    "El círculo de 400 metros es el radio con que se consultó la zona, no un centro "
    "demostrado.":
        "El círculo de 400 metros marca hasta dónde se buscó alrededor del cruce; no "
        "significa que la zona tenga ahí su centro.",
    "El círculo de 400 metros con que se consultó el cruce de Esmeralda y Paraguay.":
        "El círculo de 400 metros dentro del cual se buscó, alrededor del cruce de "
        "Esmeralda y Paraguay.",
    "El "
    "círculo es el radio con que se consultó la zona, no un centro demostrado.":
        ""
        "El círculo marca hasta dónde se buscó alrededor del cruce; no significa que la "
        "zona tenga ahí su centro.",

    # -- R19, R20, R21 y R22 · el detalle de abiertos y cerrados -------------------------
    #
    # «Al menos 254 locales» y debajo «242 abiertos y 12 cerrados», que suman 254 exactos:
    # leído de corrido, el detalle parecía desmentir el «al menos». Se dice de quiénes es
    # el detalle —de los que se relevaron— y el «al menos» queda en pie.

    "Al momento del relevamiento, 204 estaban abiertos y 7 cerrados de forma temporaria.":
        "De los locales relevados, 204 estaban abiertos y 7 cerrados de forma temporaria "
        "al momento del relevamiento.",
    "Al menos 211 locales. Al momento del relevamiento, 204 estaban abiertos y 7 cerrados "
    "de forma temporaria.":
        "Al menos 211 locales. De los relevados, 204 estaban abiertos y 7 cerrados de "
        "forma temporaria al momento del relevamiento.",
    "Al momento del relevamiento, 39 estaban abiertos y 1 cerrado de forma temporaria.":
        "De los locales relevados, 39 estaban abiertos y 1 cerrado de forma temporaria al "
        "momento del relevamiento.",
    "Al momento del relevamiento, 242 estaban abiertos y 12 cerrados de forma temporaria.":
        "De los locales relevados, 242 estaban abiertos y 12 cerrados de forma temporaria "
        "al momento del relevamiento.",
    "Al menos 254 locales. Al momento del relevamiento, 242 estaban abiertos y 12 "
    "cerrados de forma temporaria.":
        "Al menos 254 locales. De los relevados, 242 estaban abiertos y 12 cerrados de "
        "forma temporaria al momento del relevamiento.",
    "Al momento del relevamiento, 152 estaban abiertos y 6 cerrados de forma temporaria.":
        "De los locales relevados, 152 estaban abiertos y 6 cerrados de forma temporaria "
        "al momento del relevamiento.",

    # -- R20 · García del Río ------------------------------------------------------------
    #
    # La zona está rotulada «Eje o corredor» arriba de la página y tres líneas más abajo
    # decía que no alcanza para hablar de un eje ni de un corredor. La página 5 ya explica
    # por qué las dos cosas conviven: la familia describe la forma, no el tamaño.

    "No alcanza para hablar de un corredor, un eje ni un polo.":
        "No alcanza para hablar de un corredor, un eje ni un polo consolidados: la "
        "familia describe cómo se ordena la oferta, no su tamaño ni su importancia.",

    # -- R21 · La Paternal ---------------------------------------------------------------

    "No hay un centro, un corredor ni una red de puntos demostrados.":
        "No se encontró un centro, un corredor ni un conjunto de puntos vinculados entre sí.",

    # -- R22 · Villa Pueyrredón ----------------------------------------------------------

    "El contorno del mapa es el marco con que se relevó el barrio, no un área gastronómica.":
        "El contorno del mapa es el recuadro dentro del cual se buscó en el barrio, no un "
        "área gastronómica.",
    "El marco con que se relevó Villa Pueyrredón, con la oferta repartida por todo el barrio.":
        "El recuadro dentro del cual se buscó en Villa Pueyrredón, con la oferta "
        "repartida por todo el barrio.",
    "El "
    "contorno es el marco con que se relevó el barrio, no un área gastronómica: la mayor "
    "densidad del centro y centro este es una descripción, no un núcleo adoptado.":
        ""
        "Es el recuadro dentro del cual se buscó en el barrio, no un área gastronómica: "
        "la mayor densidad del centro y centro este es una descripción, no un núcleo "
        "adoptado.",
}

# original -> veces que se reemplazó. Lo llena `llano` y lo audita el QA.
APLICADAS: dict[str, int] = {}


def llano(texto: str) -> str:
    """Devuelve la frase reescrita, si esa frase exacta está en la tabla."""
    nuevo = REESCRITURAS.get(texto)
    if nuevo is None:
        return texto
    APLICADAS[texto] = APLICADAS.get(texto, 0) + 1
    return nuevo


def preparar_atlas() -> dict:
    """Deja `build_atlas_v2` configurado en la edición de conducción, sin producir nada."""
    cfg = B.CONTENT_DIR / "config_fuentes.json"
    antes = cfg.read_bytes() if cfg.is_file() else None
    _, canonical = B.validate_inputs()
    B.resolve_fonts(None)
    if antes is not None and cfg.read_bytes() != antes:
        cfg.write_bytes(antes)
    B.configurar_edicion("conduccion")
    # P-AV2-01: este generador no publica insumos, los consume. Sin `persist=False`
    # dejaba el contenido de conducción escrito en el destino de la edición técnica.
    content = B.build_public_content(canonical, persist=False)
    B.CONTENT_CACHE.clear()
    B.CONTENT_CACHE.update(content)
    B.CARTOGRAFO = CV.Cartografo(B.REPO, B.CAPAS_DIR)
    B.EXTRAS_DETALLE.clear()
    B.EXTRAS_DETALLE.update(B.cargar_extras_detalle())
    return content


def leyenda_de(rid: str) -> list[str]:
    """Los mismos elementos que la leyenda gráfica del PDF anuncia para ese mapa."""
    items = [CV.LEYENDA_AREA]
    if len(B.CARTOGRAFO.comp[B.CARTOGRAFO.comp.referencia_id == rid]) > 1:
        items.append(CV.LEYENDA_COMPONENTE)
    items.append(CV.LEYENDA_CALLE)
    return items


def pie_de_mapa(rid: str, trio: dict, descargo: str) -> list:
    """Las tres líneas rotuladas, la leyenda y el descargo: lo mismo que el PDF, en texto."""
    bloques: list = []
    for etiqueta, texto in (("Qué muestra este mapa", trio["muestra"]),
                            ("Qué mide la cifra", trio["mide"]),
                            ("Qué no es", trio["no_es"])):
        bloques.append(P([Run(etiqueta + "  ", bold=True, color=AZUL), Run(texto)], "Pie"))
    leyenda = " · ".join(leyenda_de(rid))
    bloques.append(P([Run("Leyenda  ", bold=True, color=AZUL), Run(leyenda)], "Pie"))
    if B.CARTOGRAFO.fila(rid).trazo == "punteado":
        bloques.append(txt(CV.NOTA_TRAZO_PUNTEADO, "Pie"))
    if descargo:
        bloques.append(txt(descargo, "Sec"))
    return bloques


def construir_documento(content: dict, mapas_dir: Path, dpi: int,
                        tmp: Path) -> tuple[list, list[dict], list]:
    """Arma el modelo del documento y renderiza, de paso, los 32 mapas que necesita.

    Devuelve también las «marcas»: los bloques que el índice tiene que listar, con su
    nivel. Son los ocho títulos de sección y las 22 zonas; los títulos de los 31 mapas no
    entran, porque el mapa de una zona va siempre detrás de su ficha.
    """
    bloques: list = []
    catalogo: list[dict] = []
    marcas: list[tuple[P, int]] = []
    fichas = {f["referencia_id"]: f for f in content["fichas"]}
    nombres = {rid: f["nombre"] for rid, f in fichas.items()}

    def marcado(texto: str, style: str, nivel: int) -> P:
        bloque = txt(texto, style)
        marcas.append((bloque, nivel))
        return bloque

    def mapa(clave: str, tipo: str, h_max_cm: float) -> Img:
        dx, dy = proporcion_geografica(dibujante(clave, tipo), tmp)
        w_cm, h_cm = encuadre(dx, dy, COL_W_CM, h_max_cm)
        destino = mapas_dir / (clave if clave.endswith(".png") else f"{clave}.png")
        catalogo.append(render_mapa(clave, tipo, w_cm, h_cm, dpi, destino, tmp))
        return Img(destino, w_cm, h_cm)

    # -- 1. Portada ---------------------------------------------------------------------
    bloques += [
        txt("DGDGAS", "Portada.marca"),
        txt("Dirección General de Desarrollo Gastronómico", "Portada.org"),
        txt("Atlas de referencias gastronómicas de la Ciudad de Buenos Aires",
            "Portada.titulo"),
        txt(SUBTITULO_PORTADA, "Portada.sub"),
        txt("Dónde hay oferta gastronómica en la Ciudad y cómo es cada lugar, según el "
            "trabajo de la Dirección General de Desarrollo Gastronómico.", "Portada.cuerpo"),
        txt("Qué es este Atlas", "Portada.rotulo"),
        txt("Son veintidós zonas distintas entre sí, reunidas en un mismo documento. Ni "
            "sus categorías ni sus cifras forman un ranking, y no son veintidós polos "
            "equivalentes.", "Sec"),
        txt("Los datos", "Portada.rotulo"),
        txt(C.FUENTE_UNICA, "Sec"),
        txt(NOTA_EDITABLE, "Portada.nota"),
        # Corrección editorial 2026-08-05: acá iba «Los mapas sirven para ubicar la oferta:
        # no son planos oficiales.», de la misma familia que las cuatro fórmulas retiradas y
        # exclusiva de la portada de la editable. No estaba en el Doc ni en la versión
        # original. `NOTA_EDITABLE`, en cambio, se queda: es propia de esta edición.
        Salto(),
        MarcaIndice(),
    ]

    # -- 2. Resumen ejecutivo -----------------------------------------------------------
    cantidad, lista = B.comunas_resumen()
    familias: dict[str, int] = {}
    for rid in B.EXPECTED_REFS:
        familias[B.FAMILIA_DE[rid]] = familias.get(B.FAMILIA_DE[rid], 0) + 1
    resumen_fam = "; ".join(f"{B.FAMILIA_ETIQUETA[k].lower()}, {v}"
                            for k, v in sorted(familias.items(), key=lambda kv: -kv[1]))
    bloques += [
        marcado("Resumen ejecutivo", "H1", 1),
        # Corrección editorial 2026-08-04: el resumen abre diciendo qué es un atlas y para
        # qué es éste, antes de la primera cifra.
        txt(C.QUE_ES_UN_ATLAS, "Normal"),
        txt(C.ATLAS_REUNE, "Normal"),
        txt(C.CAJA_QUE_MUESTRA_TITULO, "H2"),
        # La definición de familias es la misma corrección que `render_conduccion`: las dos
        # son la línea de conducción y tienen que decir lo mismo.
        txt(f"La Dirección identificó y caracterizó 22 zonas gastronómicas de la Ciudad, "
            f"presentes en {cantidad} de las 15 comunas (comunas {lista}). Se agrupan en "
            f"cinco familias territoriales, que describen la forma en que la oferta se ordena "
            f"en el terreno: {resumen_fam}. No se identificaron zonas en el "
            f"extremo sur de la Ciudad.", "Normal"),
        txt(C.TITULO_GRUPOS, "H2"),
        txt(C.NOTA_GRUPOS, "Sec"),
    ]
    for titulo, datos, nota in C.GRUPOS_RESUMEN:
        bloques += [txt(titulo, "H3"), txt(datos, "Normal"), txt(nota, "SecIt")]
    bloques += [
        txt(C.CAJA_USO_TITULO, "H2"),
        txt(C.CAJA_USO, "Normal"),
        txt(C.FUENTE_UNICA, "Sec"),
        Salto(),
    ]

    # -- 3. Cómo se construyeron las zonas (sección nueva) ------------------------------
    bloques.append(marcado(COMO_SE_CONSTRUYERON_TITULO, "H1", 1))
    for clase, valor in COMO_SE_CONSTRUYERON:
        if clase == "h":
            bloques.append(txt(valor, "H2"))
        elif clase == "p":
            bloques.append(txt(valor, "Normal"))
        else:
            bloques.append(Bul([[Run(x)] for x in valor]))
    bloques.append(Salto())

    # -- 4. Cómo leer el Atlas ----------------------------------------------------------
    bloques.append(marcado("Cómo leer el Atlas", "H1", 1))
    for titulo, texto in C.COMO_LEER:
        bloques += [txt(titulo, "H2"), txt(texto, "Normal")]
    bloques += [txt(C.CAJA_CIERRE_TITULO, "H2"), txt(CIERRE_COMO_LEER, "Normal"), Salto()]

    # -- 5. Las cinco familias territoriales --------------------------------------------
    bloques += [
        marcado("Las cinco familias territoriales", "H1", 1),
        txt(SUBTITULO_FAMILIAS, "Sec"),
    ]
    for fam, etiqueta in B.FAMILIA_ETIQUETA.items():
        refs = [r for r in B.EXPECTED_REFS if B.FAMILIA_DE[r] == fam]
        bloques += [
            txt(f"{etiqueta} ({len(refs)})", "H2"),
            txt(DEFINICION_FAMILIA[fam], "Sec"),
            Bul([[Run(f"{r} · ", bold=True, color=AZUL), Run(nombres[r])] for r in refs]),
        ]
    bloques += [txt(C.CAJA_FAMILIAS_TITULO, "H2"), txt(C.CAJA_FAMILIAS, "Normal"), Salto()]

    # -- 6. Las 22 zonas en la Ciudad ---------------------------------------------------
    cabecera_general = [
        marcado("Las 22 zonas en la Ciudad", "H1", 1),
        txt(SUBTITULO_MAPA_GENERAL, "Sec"),
    ]
    familias_runs: list[Run] = [Run("Familias  ", bold=True, color=AZUL)]
    for i, (fam, etiqueta) in enumerate(B.FAMILIA_ETIQUETA.items()):
        if i:
            familias_runs.append(Run("   ·   ", color=GRIS))
        familias_runs.append(Run(etiqueta, bold=True, color=FAMILIA_COLOR[fam].lstrip("#")))
    pie_general = [P(familias_runs, "Pie")]
    for texto, fnt in B.notas_locator(True):
        pie_general.append(txt(texto, "Warn" if fnt == "AtlasSans-Bold" else "Pie"))
    alto_fijo = sum(alto_bloque(b) for b in cabecera_general + pie_general)
    h_max = (COL_H - alto_fijo - HOLGURA_CM * CM
             - STYLES["Imagen"]["before"] - STYLES["Imagen"]["after"]) / CM
    bloques += cabecera_general + [mapa("MAPA_GENERAL", "general", h_max)] + pie_general
    bloques.append(Salto())

    # -- 7. Las 22 zonas, una por una ---------------------------------------------------
    bloques.append(marcado(SECCION_ZONAS_TITULO, "H1", 1))
    for pos, rid in enumerate(B.EXPECTED_REFS):
        ficha = fichas[rid]
        fuente_ficha = C.FICHAS[rid]
        if pos:
            bloques.append(Salto())
        bloques += [
            marcado(f"{rid} · {ficha['nombre']}", "H2", 2),
            txt(B.FAMILIA_ETIQUETA[B.FAMILIA_DE[rid]], "Sec"),
            txt("Qué es esta zona", "Bloque"),
            txt(fuente_ficha["es"], "Normal"),
            txt("Cuánta oferta hay", "Bloque"),
            P([Run(C.frase_cifra(ficha), bold=True, color=AZUL)], "Normal"),
            txt(fuente_ficha["contexto"], "Sec"),
            txt("De qué se compone", "Bloque"),
            Bul([[Run(x)] for x in fuente_ficha["compone"]]),
            txt("Qué hay que tener en cuenta", "Bloque"),
            Bul([[Run(x)] for x in fuente_ficha["cuenta"]], "BulletSec"),
        ]
        # Mapa principal, en su propia página.
        titulo_mapa = txt(f"{rid} · {ficha['nombre']} · Mapa", "H3", keep_with_next=True)
        pie = pie_de_mapa(rid, C.TRIOS[rid], B.DESCARGO_MAPA["conduccion"])
        alto_fijo = alto_bloque(titulo_mapa) + sum(alto_bloque(b) for b in pie)
        h_max = (COL_H - alto_fijo - HOLGURA_CM * CM
                 - STYLES["Imagen"]["before"] - STYLES["Imagen"]["after"]) / CM
        bloques += [Salto(), titulo_mapa, mapa(rid, "principal", h_max)] + pie
        # Mapas ampliados de la zona, uno por página.
        for archivo in B.AMPLIACIONES_DE.get(rid, []):
            clave = archivo
            titulo = txt(f"{rid} · {C.TITULOS_DETALLE[archivo]} · Más de cerca", "H3",
                         keep_with_next=True)
            pie = pie_de_mapa(rid, C.TRIOS_DETALLE[archivo],
                              B.DESCARGO_AMPLIACION["conduccion"])
            alto_fijo = alto_bloque(titulo) + sum(alto_bloque(b) for b in pie)
            h_max = (COL_H - alto_fijo - HOLGURA_CM * CM
                     - STYLES["Imagen"]["before"] - STYLES["Imagen"]["after"]) / CM
            bloques += [Salto(), titulo, mapa(clave, "ampliacion", h_max)] + pie
    bloques.append(Salto())

    # -- 8. Las 22 zonas de un vistazo --------------------------------------------------
    filas = [[rid, fichas[rid]["nombre"], B.FAMILIA_ETIQUETA[B.FAMILIA_DE[rid]],
              C.frase_cifra(fichas[rid])] for rid in B.EXPECTED_REFS]
    bloques += [
        marcado("Las 22 zonas de un vistazo", "H1", 1),
        # C-03 · El .docx no tiene índice ni números de página, así que no puede remitir a
        # nada: la promesa se saca, igual que se sacó del subtítulo del mapa general.
        txt("Tabla de consulta rápida.", "Sec"),
        Tbl(C.ANEXO_A_ENCABEZADOS, filas, [1.5, 4.5, 4.2, 6.8]),
        txt(C.ANEXO_A_PIE, "Sec"),
        Salto(),
    ]

    # -- 9. Cómo se hizo este Atlas y qué no dice ---------------------------------------
    bloques.append(marcado(C.COMO_SE_HIZO_TITULO, "H1", 1))
    for parrafo in C.COMO_SE_HIZO:
        bloques.append(txt(parrafo, "Normal"))

    return bloques, catalogo, marcas


# =======================================================================================
# Paginación: el modelo se corta en páginas antes de escribir nada
# =======================================================================================


# Estilos que en Word llevan «conservar con el siguiente»: un título nunca queda solo al
# pie de una página. La medición tiene que respetarlo o el .pdf de prueba y el .docx se
# separan justo donde más importa.
ESTILOS_PEGADOS = {"H1", "H2", "H3", "Bloque", "Portada.titulo", "Portada.rotulo",
                   "Portada.marca"}


def pegado_al_siguiente(bloque) -> bool:
    return isinstance(bloque, P) and (bloque.keep_with_next or bloque.style in ESTILOS_PEGADOS)


def agrupar(bloques: list) -> list[list]:
    """Junta cada título con lo que lo sigue: se paginan como una sola pieza."""
    grupos: list[list] = []
    actual: list = []
    for bloque in bloques:
        actual.append(bloque)
        if not pegado_al_siguiente(bloque):
            grupos.append(actual)
            actual = []
    if actual:
        grupos.append(actual)
    return grupos


INDICE_TITULO = "Índice"
INDICE_NOTA = ("Cada zona aparece en la página donde empieza su ficha; su mapa viene "
               "inmediatamente después.")


def armar_indice(bloques: list, marcas: list) -> tuple[list, list[list]]:
    """Reemplaza la marca del índice por el índice, con la página real de cada entrada.

    Es un punto fijo: el índice ocupa páginas y esas páginas corren a todas las demás.
    Se repite hasta que los folios dejan de moverse —con un índice de una hoja converge
    en dos vueltas— y si no convergiera corta con error, antes que publicar un índice que
    manda a la página equivocada.
    """
    folios: dict[int, int] = {id(b): 0 for b, _ in marcas}
    for _ in range(8):
        armado: list = []
        for bloque in bloques:
            if not isinstance(bloque, MarcaIndice):
                armado.append(bloque)
                continue
            armado.append(txt(INDICE_TITULO, "H1"))
            armado.append(txt(INDICE_NOTA, "Sec"))
            for marca, nivel in marcas:
                armado.append(Idx(marca.runs[0].text, folios[id(marca)],
                                  "Indice" if nivel == 1 else "IndiceZona"))
            armado.append(Salto())
        paginas = paginar(armado)
        ubicacion = {id(b): n for n, pagina in enumerate(paginas, start=1) for b in pagina}
        nuevos = {id(b): ubicacion[id(b)] for b, _ in marcas}
        if nuevos == folios:
            return armado, paginas
        folios = nuevos
    raise RuntimeError("el índice no converge: revisar la paginación")


def paginar(bloques: list) -> list[list]:
    """Reparte los bloques en páginas.

    Los saltos son explícitos: cada página del documento está delimitada por uno. Esto
    sólo verifica que lo que va entre dos saltos entre en la caja de texto, y corta si no
    (sólo puede pasar en la tabla del anexo, la única corrida larga sin salto propio).
    """
    paginas: list[list] = [[]]
    alto = 0.0

    def nueva_pagina() -> None:
        nonlocal alto
        paginas.append([])
        alto = 0.0

    for grupo in agrupar(bloques):
        for bloque in grupo:
            if isinstance(bloque, Salto):
                nueva_pagina()
                continue
            if isinstance(bloque, Tbl):
                for parte in partir_tabla(bloque, COL_H - alto):
                    if paginas[-1] and alto + alto_tabla(parte) > COL_H:
                        nueva_pagina()
                    paginas[-1].append(parte)
                    alto += alto_tabla(parte)
                continue
            h_grupo = sum(alto_bloque(b) for b in grupo if not isinstance(b, (Salto, Tbl)))
            h = alto_bloque(bloque)
            # El grupo entero se muda si no entra; si ni siquiera cabe en una página
            # vacía, se pagina bloque a bloque para no entrar en un bucle.
            necesita = h_grupo if bloque is grupo[0] and h_grupo <= COL_H else h
            if paginas[-1] and alto + necesita > COL_H:
                nueva_pagina()
            paginas[-1].append(bloque)
            alto += h
    return [p for p in paginas if p]


def partir_tabla(t: Tbl, disponible: float) -> list[Tbl]:
    """Corta la tabla por filas enteras; la cabecera se repite en cada trozo."""
    anchos = [w * CM for w in t.widths_cm]
    h_cab = alto_fila(t.headers, anchos, "TablaCab") + 2
    partes: list[Tbl] = []
    actual: list[list[str]] = []
    alto = h_cab
    tope = disponible
    for fila in t.rows:
        h = alto_fila(fila, anchos, "Tabla")
        if actual and alto + h > tope:
            partes.append(Tbl(t.headers, actual, t.widths_cm))
            actual, alto, tope = [], h_cab, COL_H
        actual.append(fila)
        alto += h
    if actual:
        partes.append(Tbl(t.headers, actual, t.widths_cm))
    return partes


# =======================================================================================
# Emisión del .docx
# =======================================================================================


def _rfonts(elemento, nombre: str = "Calibri") -> None:
    rpr = elemento.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    # Los estilos de título de la plantilla apuntan a la fuente del tema. Se quitan: el
    # documento tiene que verse igual en una máquina que no traiga ese tema.
    for attr in ("w:asciiTheme", "w:hAnsiTheme", "w:eastAsiaTheme", "w:cstheme"):
        if rfonts.get(qn(attr)) is not None:
            del rfonts.attrib[qn(attr)]
    for attr in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
        rfonts.set(qn(attr), nombre)


def _idioma(elemento, valor: str = "es-AR") -> None:
    rpr = elemento.get_or_add_rPr()
    lang = rpr.find(qn("w:lang"))
    if lang is None:
        lang = OxmlElement("w:lang")
        rpr.append(lang)
    lang.set(qn("w:val"), valor)
    lang.set(qn("w:bidi"), "ar-SA")


def preparar_estilos(doc: docx.Document) -> None:
    """Deja los estilos de Word con la escala del documento. Nada de formato manual."""
    existentes = {s.name for s in doc.styles}
    creados = {
        "Portada marca": "Normal", "Portada organismo": "Normal",
        "Portada rotulo": "Normal", "Texto secundario": "Normal",
        "Texto secundario cursiva": "Texto secundario", "Advertencia": "Normal",
        "Bloque de ficha": "Normal", "Pie de mapa": "Normal",
        "Imagen de mapa": "Normal", "Celda de tabla": "Normal",
        "Cabecera de tabla": "Celda de tabla", "List Bullet Sec": "List Bullet",
        "Indice seccion": "Normal", "Indice zona": "Normal",
    }
    for nombre, base in creados.items():
        if nombre not in existentes:
            estilo = doc.styles.add_style(nombre, WD_STYLE_TYPE.PARAGRAPH)
            estilo.base_style = doc.styles[base]
            estilo.quick_style = True

    aplicados: set[str] = set()
    for clave, st in STYLES.items():
        nombre = st["word"]
        if nombre in aplicados:
            continue
        aplicados.add(nombre)
        estilo = doc.styles[nombre]
        estilo.font.name = "Calibri"
        estilo.font.size = Pt(st["size"])
        estilo.font.bold = bool(st.get("bold", False))
        estilo.font.italic = bool(st.get("italic", False))
        estilo.font.color.rgb = RGBColor.from_string(st["color"])
        _rfonts(estilo.element)
        _idioma(estilo.element)
        pf = estilo.paragraph_format
        pf.space_before = Pt(st["before"])
        pf.space_after = Pt(st["after"])
        pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        pf.line_spacing = st["mult"]
        pf.alignment = WD_ALIGN_PARAGRAPH.LEFT
        pf.widow_control = True
        if nombre.startswith("Heading") or nombre in ("Title", "Bloque de ficha"):
            pf.keep_with_next = True

    for nombre in ("List Bullet", "List Bullet Sec"):
        pf = doc.styles[nombre].paragraph_format
        pf.left_indent = Cm(SANGRIA_BULLET_CM)
        pf.first_line_indent = Cm(-SANGRIA_BULLET_CM * 0.55)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    _rfonts(normal.element)
    _idioma(normal.element)


def campo_pagina(parrafo) -> None:
    """Escribe { PAGE } con su resultado en caché, tal como lo guarda Word.

    Los tres `w:fldChar` —begin, separate, end— con el `w:instrText` en el medio son la
    forma larga del campo. Se usa esa y no `w:fldSimple` porque es la que escribe el
    propio Word, y porque el resultado en caché deja un número visible aunque el lector no
    actualice los campos. Word los recalcula al abrir y al imprimir.
    """
    def nuevo_run():
        r = OxmlElement("w:r")
        parrafo._p.append(r)
        return r

    inicio = OxmlElement("w:fldChar")
    inicio.set(qn("w:fldCharType"), "begin")
    nuevo_run().append(inicio)

    instruccion = OxmlElement("w:instrText")
    instruccion.set(qn("xml:space"), "preserve")
    instruccion.text = " PAGE "
    nuevo_run().append(instruccion)

    separador = OxmlElement("w:fldChar")
    separador.set(qn("w:fldCharType"), "separate")
    nuevo_run().append(separador)

    resultado = OxmlElement("w:t")
    resultado.text = "1"
    nuevo_run().append(resultado)

    fin = OxmlElement("w:fldChar")
    fin.set(qn("w:fldCharType"), "end")
    nuevo_run().append(fin)


def poner_folios(doc: docx.Document) -> None:
    """Número de página abajo a la derecha, en todas las páginas menos la portada."""
    seccion = doc.sections[0]
    seccion.different_first_page_header_footer = True
    seccion.first_page_footer.is_linked_to_previous = False      # la portada, sin folio
    pie = seccion.footer
    pie.is_linked_to_previous = False
    parrafo = pie.paragraphs[0]
    parrafo.style = doc.styles[STYLES["PiePagina"]["word"]]
    parrafo.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    campo_pagina(parrafo)


def escribir_indice(doc, bloque: Idx):
    """Una línea del índice: título, guía de puntos y folio contra el margen derecho."""
    st = STYLES[bloque.style]
    p = doc.add_paragraph(style=st["word"])
    p.paragraph_format.tab_stops.add_tab_stop(
        Cm(COL_W_CM), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.DOTS)
    p.add_run(bloque.texto)
    p.add_run("\t" + str(bloque.pagina))
    return p


def escribir_parrafo(doc, bloque: P):
    st = STYLES[bloque.style]
    p = doc.add_paragraph(style=st["word"])
    for run in bloque.runs:
        r = p.add_run(run.text)
        if run.bold:
            r.bold = True
        if run.italic:
            r.italic = True
        if run.color:
            r.font.color.rgb = RGBColor.from_string(run.color)
    if bloque.keep_with_next:
        p.paragraph_format.keep_with_next = True
    return p


def _borde(elemento, lado: str, valor: str, color: str = GRIS_LINEA, sz: str = "4") -> None:
    b = OxmlElement(f"w:{lado}")
    b.set(qn("w:val"), valor)
    if valor != "none":
        b.set(qn("w:sz"), sz)
        b.set(qn("w:space"), "0")
        b.set(qn("w:color"), color)
    elemento.append(b)


def escribir_tabla(doc, t: Tbl) -> None:
    tabla = doc.add_table(rows=1 + len(t.rows), cols=len(t.headers))
    tabla.alignment = WD_TABLE_ALIGNMENT.LEFT
    tabla.autofit = False        # deja `w:tblLayout` en «fixed»

    # `w:tblPr` es una secuencia, no una bolsa: tblW · jc · tblBorders · tblLayout · tblLook.
    # Agregar los elementos al final deja el XML fuera de esquema y Word abre el archivo
    # avisando de «contenido ilegible». Se reusa lo que ya puso python-docx y el bloque de
    # bordes se inserta en su lugar.
    tbl_pr = tabla._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.insert(0, tbl_w)
    tbl_w.set(qn("w:type"), "dxa")
    tbl_w.set(qn("w:w"), str(int(round(sum(t.widths_cm) * 567))))

    # Sólo filetes horizontales: sin cajas, sin verticales, sin relleno de color.
    bordes = OxmlElement("w:tblBorders")
    for lado in ("top", "left", "bottom", "right", "insideH", "insideV"):
        _borde(bordes, lado, "none" if lado in ("left", "right", "insideV") else "single")
    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is not None:
        layout.addprevious(bordes)
    else:
        tbl_pr.append(bordes)

    for i, w in enumerate(t.widths_cm):
        tabla.columns[i].width = Cm(w)

    for j, (celda, texto) in enumerate(zip(tabla.rows[0].cells, t.headers)):
        celda.width = Cm(t.widths_cm[j])
        celda.paragraphs[0].style = doc.styles["Cabecera de tabla"]
        celda.paragraphs[0].add_run(texto)
    tr_pr = tabla.rows[0]._tr.get_or_add_trPr()
    tr_pr.append(OxmlElement("w:tblHeader"))

    for fila, valores in zip(tabla.rows[1:], t.rows):
        fila._tr.get_or_add_trPr().append(OxmlElement("w:cantSplit"))
        for j, (celda, texto) in enumerate(zip(fila.cells, valores)):
            celda.width = Cm(t.widths_cm[j])
            celda.paragraphs[0].style = doc.styles["Celda de tabla"]
            celda.paragraphs[0].add_run(texto)


def corregir_settings(doc: docx.Document) -> None:
    """C-02 · `w:zoom` sin `w:percent`.

    La plantilla que trae `python-docx` escribe `<w:zoom w:val="bestFit"/>`. En el esquema
    de OOXML, `CT_Zoom` declara `w:val` opcional y `w:percent` **requerido**, así que ese
    elemento es inválido. Word lo tolera en la práctica, pero era la única falla de
    validación del archivo y no hay Word en esta máquina para comprobar qué hace con él.
    Se completa con el 100 %, que es el valor neutro: `w:val="bestFit"` sigue mandando.
    """
    zoom = doc.settings.element.find(qn("w:zoom"))
    if zoom is not None and zoom.get(qn("w:percent")) is None:
        zoom.set(qn("w:percent"), "100")


def emitir_docx(paginas: list[list], destino: Path) -> None:
    doc = docx.Document()
    corregir_settings(doc)
    seccion = doc.sections[0]
    seccion.page_width, seccion.page_height = Cm(PAGE_W_CM), Cm(PAGE_H_CM)
    for lado in ("top", "bottom", "left", "right"):
        setattr(seccion, f"{lado}_margin", Cm(MARGIN_CM))
    seccion.header_distance = Cm(1.2)
    seccion.footer_distance = Cm(1.2)
    preparar_estilos(doc)
    poner_folios(doc)

    propiedades = doc.core_properties
    propiedades.title = "Atlas de referencias gastronómicas de la Ciudad de Buenos Aires"
    propiedades.author = "DGDGAS - Dirección General de Desarrollo Gastronómico"
    propiedades.category = "Edición de conducción · documento editable"
    propiedades.comments = ("Edición editable de la edición de conducción del Atlas. "
                            "Los mapas sirven para ubicar la oferta: no son planos oficiales.")

    primero = True
    for pagina in paginas:
        if not primero:
            salto = doc.add_paragraph()
            salto.paragraph_format.space_before = Pt(0)
            salto.paragraph_format.space_after = Pt(0)
            salto.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
            salto.paragraph_format.line_spacing = Pt(1)
            salto.add_run().add_break(WD_BREAK.PAGE)
        primero = False
        for bloque in pagina:
            if isinstance(bloque, P):
                escribir_parrafo(doc, bloque)
            elif isinstance(bloque, Bul):
                for item in bloque.items:
                    escribir_parrafo(doc, P(item, bloque.style))
            elif isinstance(bloque, Img):
                st = STYLES["Imagen"]
                p = doc.add_paragraph(style=st["word"])
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.paragraph_format.keep_together = True
                p.add_run().add_picture(str(bloque.path),
                                        width=Cm(bloque.w_cm), height=Cm(bloque.h_cm))
            elif isinstance(bloque, Tbl):
                escribir_tabla(doc, bloque)
            elif isinstance(bloque, Idx):
                escribir_indice(doc, bloque)
            else:
                raise TypeError(bloque)
    doc.save(destino)


# =======================================================================================
# Prueba de paginación en PDF: mismo modelo, misma métrica, mismas cajas
# =======================================================================================


def emitir_pdf(paginas: list[list], destino: Path) -> list[float]:
    c = rl_canvas.Canvas(str(destino), pagesize=(PAGE_W_CM * CM, PAGE_H_CM * CM))
    c.setTitle("Atlas de referencias gastronómicas de la Ciudad de Buenos Aires")
    c.setAuthor("DGDGAS - Dirección General de Desarrollo Gastronómico")
    x0 = MARGIN_CM * CM
    techo = (PAGE_H_CM - MARGIN_CM) * CM
    st_pie = STYLES["PiePagina"]
    usados: list[float] = []
    for n, pagina in enumerate(paginas, start=1):
        y = techo
        for bloque in pagina:
            y = dibujar_bloque(c, bloque, x0, y)
        usados.append((techo - y) / CM)
        if n > 1:                                   # la portada no lleva folio
            c.setFont(fuente(False, False), st_pie["size"])
            c.setFillColor(HexColor("#" + st_pie["color"]))
            c.drawRightString((PAGE_W_CM - MARGIN_CM) * CM,
                              1.2 * CM + st_pie["size"] * 0.22, str(n))
        c.showPage()
    c.save()
    return usados


def dibujar_lineas(c, lineas, x: float, y: float, st: dict) -> float:
    salto = st["size"] * CALIBRI_SINGLE * st["mult"]
    for linea in lineas:
        y -= salto
        cx = x
        for texto, bold, italic, color in linea:
            c.setFont(fuente(bold, italic), st["size"])
            c.setFillColor(HexColor("#" + color))
            c.drawString(cx, y + st["size"] * 0.22, texto)
            cx += ancho(texto, bold, italic, st["size"])
    return y


def dibujar_bloque(c, bloque, x: float, y: float) -> float:
    if isinstance(bloque, Idx):
        st = STYLES[bloque.style]
        y -= st["before"]
        lineas = partir([Run(bloque.texto)], COL_W - ANCHO_FOLIO, st)
        y_texto = dibujar_lineas(c, lineas, x, y, st)
        folio = str(bloque.pagina)
        c.setFont(fuente(st["bold"], False), st["size"])
        c.setFillColor(HexColor("#" + st["color"]))
        base = y_texto + st["size"] * 0.22
        c.drawRightString(x + COL_W, base, folio)
        # Guía de puntos entre el título y el folio, como la de un índice de Word.
        desde = x + ancho(lineas[-1][-1][0], st["bold"], False, st["size"]) + 3
        if len(lineas) == 1:
            desde = x + ancho(bloque.texto, st["bold"], False, st["size"]) + 3
        hasta = x + COL_W - ancho(folio, st["bold"], False, st["size"]) - 3
        if hasta > desde:
            c.setFillColor(HexColor("#" + GRIS))
            paso = ancho(". ", False, False, st["size"])
            n = int((hasta - desde) / paso)
            c.setFont(fuente(False, False), st["size"])
            c.drawString(desde, base, ". " * max(n, 0))
        return y_texto - st["after"]
    if isinstance(bloque, P):
        st = STYLES[bloque.style]
        y -= st["before"]
        y = dibujar_lineas(c, partir(bloque.runs, COL_W, st), x, y, st)
        return y - st["after"]
    if isinstance(bloque, Bul):
        st = STYLES[bloque.style]
        y -= st["before"]
        sangria = SANGRIA_BULLET_CM * CM
        for item in bloque.items:
            lineas = partir(item, COL_W - sangria, st)
            c.setFont(fuente(False, False), st["size"])
            c.setFillColor(HexColor("#" + st["color"]))
            c.drawString(x + sangria * 0.45,
                         y - st["size"] * CALIBRI_SINGLE * st["mult"] + st["size"] * 0.22, "•")
            y = dibujar_lineas(c, lineas, x + sangria, y, st)
        return y - st["after"]
    if isinstance(bloque, Img):
        st = STYLES["Imagen"]
        y -= st["before"]
        w, h = bloque.w_cm * CM, bloque.h_cm * CM
        c.drawImage(str(bloque.path), x + (COL_W - w) / 2, y - h, w, h,
                    preserveAspectRatio=False, anchor="c")
        return y - h - st["after"]
    if isinstance(bloque, Tbl):
        return dibujar_tabla(c, bloque, x, y)
    raise TypeError(bloque)


def dibujar_tabla(c, t: Tbl, x: float, y: float) -> float:
    anchos = [w * CM for w in t.widths_cm]
    total = sum(anchos)
    c.setStrokeColor(HexColor("#" + GRIS_LINEA))
    c.setLineWidth(0.5)
    c.line(x, y, x + total, y)
    for valores, estilo in [(t.headers, "TablaCab")] + [(f, "Tabla") for f in t.rows]:
        st = STYLES[estilo]
        h = alto_fila(valores, anchos, estilo)
        cx = x
        for valor, w in zip(valores, anchos):
            lineas = partir([Run(valor)], w - 8.0, st)
            dibujar_lineas(c, lineas, cx + 4.0, y - PAD_CELDA, st)
            cx += w
        y -= h
        c.line(x, y, x + total, y)
    return y - 6


# =======================================================================================
# QA
# =======================================================================================


def sha256(path: Path) -> str:
    h = hashlib.sha256()
    with path.open("rb") as fh:
        for trozo in iter(lambda: fh.read(1 << 20), b""):
            h.update(trozo)
    return h.hexdigest()


PROHIBIDOS_INSTITUCIONALES = ["DataGastro", "borrador", "preliminar",
                              "documento interno", "a validar"]

# Apariciones del vocabulario vigilado que la Dirección haya autorizado a dejar tal cual.
# El mecanismo se conserva —una excepción declarada se informa como OBSERVADO en vez de
# corregirse en silencio— pero hoy está vacío: la única que hubo, «absorba» en «Cómo se
# construyeron las zonas», se reemplazó por indicación de la Dirección (C-01) y el archivo
# quedó en cero términos vigilados.
EXCEPCIONES_VOCABULARIO: dict[tuple[str, str], str] = {}


# Números que cambian a propósito entre la edición de conducción y esta, con su motivo.
# No son cifras del Atlas: son referencias de navegación. Todo lo demás tiene que dar cero.
EXCEPCIONES_CIFRAS = {
    "6": "«el mapa de la página 6» pasó a «el mapa general, unas páginas más adelante»: el "
         "índice nuevo da la página exacta y no se desfasa cuando Jefatura edite el archivo.",
}

MC = "http://schemas.openxmlformats.org/markup-compatibility/2006"

# Partes que se validan contra el esquema, y las que quedan afuera con su motivo. Las dos
# excluidas llegan intactas desde la plantilla de `python-docx`: el control lo comprueba
# byte a byte, así que la exclusión no esconde nada que este generador haya escrito.
PARTES_VALIDADAS = [
    "[Content_Types].xml", "_rels/.rels", "word/_rels/document.xml.rels",
    "customXml/_rels/item1.xml.rels", "docProps/core.xml", "docProps/app.xml",
    "customXml/itemProps1.xml", "word/document.xml", "word/settings.xml",
    "word/styles.xml", "word/stylesWithEffects.xml", "word/numbering.xml",
    "word/fontTable.xml", "word/webSettings.xml", "word/footer1.xml", "word/footer2.xml",
]
PARTES_EXCLUIDAS = {
    "word/theme/theme1.xml": "tema DrawingML de la plantilla",
    "customXml/item1.xml": "XML propio, sin esquema declarado",
}


def compatibilidad(raiz):
    """Preproceso de compatibilidad (MCE): saca lo que `mc:Ignorable` declara ignorable.

    Es lo que hace un lector conforme antes de mirar el contenido, y es lo que deja el
    archivo comparable con el esquema: `w14:docId` y compañía viven en un namespace que
    `mc:Ignorable` marca como descartable y que ECMA-376 no describe.
    """
    ignorables: set[str] = set()
    for elemento in raiz.iter():
        if not isinstance(elemento.tag, str):
            continue
        declarado = elemento.get(f"{{{MC}}}Ignorable")
        for prefijo in (declarado or "").split():
            ignorables.add(elemento.nsmap.get(prefijo))
    for elemento in list(raiz.iter()):
        if not isinstance(elemento.tag, str):
            continue
        for atributo in list(elemento.attrib):
            ns = atributo[1:].partition("}")[0] if atributo.startswith("{") else None
            if ns == MC or ns in ignorables:
                del elemento.attrib[atributo]
    for elemento in list(raiz.iter()):
        if not isinstance(elemento.tag, str):
            continue
        ns = elemento.tag[1:].partition("}")[0] if elemento.tag.startswith("{") else None
        if ns in ignorables and elemento.getparent() is not None:
            elemento.getparent().remove(elemento)
    return raiz


def validar_esquema(docx_path: Path) -> tuple[list[str], int, int]:
    """Valida las partes XML del paquete contra el esquema de `qa/esquema/`.

    Devuelve los errores, cuántas partes se validaron y cuántos elementos distintos había
    en ellas: ese último número es lo que evita que el esquema pase por cubrir poco, porque
    todo elemento no declarado es un error de validación, no un silencio.
    """
    from lxml import etree

    esquema = etree.XMLSchema(etree.parse(str(B.QA_DIR / "esquema" / "catalogo.xsd")))
    errores: list[str] = []
    elementos: set[str] = set()
    with zipfile.ZipFile(docx_path) as paquete:
        presentes = set(paquete.namelist())
        for parte in PARTES_VALIDADAS:
            if parte not in presentes:
                errores.append(f"{parte}: falta en el paquete")
                continue
            raiz = compatibilidad(etree.fromstring(paquete.read(parte)))
            elementos |= {e.tag for e in raiz.iter() if isinstance(e.tag, str)}
            if not esquema.validate(etree.ElementTree(raiz)):
                primero = list(esquema.error_log)[0]
                errores.append(f"{parte}:{primero.line}: {primero.message}")
        sin_declarar = [p for p in presentes
                        if (p.endswith(".xml") or p.endswith(".rels"))
                        and p not in PARTES_VALIDADAS and p not in PARTES_EXCLUIDAS]
        if sin_declarar:
            errores.append("partes fuera del alcance declarado: " + ", ".join(sorted(sin_declarar)))
        # Las dos exclusiones sólo valen si de verdad son las de la plantilla.
        plantilla = Path(docx.__file__).resolve().parent / "templates" / "default.docx"
        with zipfile.ZipFile(plantilla) as origen:
            for parte in PARTES_EXCLUIDAS:
                if paquete.read(parte) != origen.read(parte):
                    errores.append(f"{parte}: excluida por venir de la plantilla, pero difiere")
    return errores, len(PARTES_VALIDADAS), len(elementos)


NUMERO = re.compile(r"\d+(?:[.,]\d+)*")


def qa(docx_path: Path, pdf_path: Path, paginas: list[list], usados: list[float],
       catalogo: list[dict], destino: Path) -> dict:
    filas: list[dict] = []

    def check(control: str, ok: bool, detalle: str, estado: str | None = None) -> None:
        filas.append({"control": control,
                      "resultado": estado or ("PASS" if ok else "FAIL"),
                      "detalle": detalle})

    # 1 · el .docx abre y tiene la estructura declarada
    d = docx.Document(docx_path)
    estilos_titulo = [p.style.name for p in d.paragraphs
                      if p.style.name in ("Title", "Heading 1", "Heading 2", "Heading 3")]
    check("estilos_de_titulo_aplicados", len(estilos_titulo) > 60,
          f"{estilos_titulo.count('Heading 1')} Título 1, "
          f"{estilos_titulo.count('Heading 2')} Título 2, "
          f"{estilos_titulo.count('Heading 3')} Título 3, "
          f"{estilos_titulo.count('Title')} Título")
    imagenes = len(d.inline_shapes)
    check("31_mapas_de_zona_mas_el_general", imagenes == 32, f"{imagenes} imágenes")
    check("tabla_real_de_word", len(d.tables) >= 1,
          f"{len(d.tables)} tabla(s), {sum(len(t.rows) - 1 for t in d.tables)} filas de datos")

    texto_docx = "\n".join(p.text for p in d.paragraphs)
    texto_docx += "\n" + "\n".join(cel.text for t in d.tables for f in t.rows for cel in f.cells)

    # 2 · las 22 fichas, completas
    faltan = []
    for rid in B.EXPECTED_REFS:
        nombre = B.CONTENT_CACHE and {f["referencia_id"]: f["nombre"]
                                      for f in B.CONTENT_CACHE["fichas"]}[rid]
        if f"{rid} · {nombre}" not in texto_docx:
            faltan.append(rid)
    check("22_fichas_presentes", not faltan, f"faltan: {faltan}" if faltan else "22/22")
    bloques_ficha = [texto_docx.count(x) for x in
                     ("Qué es esta zona", "Cuánta oferta hay", "De qué se compone",
                      "Qué hay que tener en cuenta")]
    check("cuatro_bloques_por_ficha", all(n >= 22 for n in bloques_ficha),
          f"apariciones: {bloques_ficha}")
    trios = [texto_docx.count(x) for x in
             ("Qué muestra este mapa", "Qué mide la cifra", "Qué no es")]
    check("pie_de_tres_lineas_en_los_31_mapas", all(n == 31 for n in trios),
          f"apariciones: {trios}")

    # 3 · lenguaje
    hallazgos = LENG.hallazgos(texto_docx)
    declarados = [(t, f) for t, f in hallazgos
                  if (t, f.lower()) in EXCEPCIONES_VOCABULARIO]
    sin_declarar = [(t, f) for t, f in hallazgos if (t, f.lower()) not in EXCEPCIONES_VOCABULARIO]
    if sin_declarar:
        check("vocabulario_conduccion", False,
              "; ".join(f"{t} → «{f}»" for t, f in sin_declarar))
    elif declarados:
        check("vocabulario_conduccion", True,
              "; ".join(f"{t} → «{f}»: {EXCEPCIONES_VOCABULARIO[(t, f.lower())]}"
                        for t, f in declarados),
              estado="OBSERVADO")
    else:
        check("vocabulario_conduccion", True,
              f"sin hallazgos sobre {len(LENG.VOCABULARIO_PROHIBIDO)} términos vigilados")
    encontrados = [w for w in PROHIBIDOS_INSTITUCIONALES if w.lower() in texto_docx.lower()]
    check("sin_terminos_institucionales_vedados", not encontrados,
          ", ".join(encontrados) or "ninguno")
    check("marca_DGDGAS",
          "Dirección General de Desarrollo Gastronómico" in texto_docx
          and "de Gastronomía" not in texto_docx, "denominación oficial correcta")

    # 4 · paginación: ninguna página se pasa de la caja
    excedidas = [(i + 1, round(u, 2)) for i, u in enumerate(usados) if u > COL_H_CM]
    check("ninguna_pagina_desborda", not excedidas, f"máximo {max(usados):.2f} cm de {COL_H_CM}")
    holgura_minima = min(COL_H_CM - u for u in usados)
    check("holgura_minima_por_pagina", holgura_minima >= 0.4,
          f"{holgura_minima:.2f} cm en la página más llena")

    # 5 · imágenes: entera, en una sola página, y con resolución suficiente
    problemas = []
    for pagina in paginas:
        for bloque in pagina:
            if not isinstance(bloque, Img):
                continue
            fijo = sum(alto_bloque(b) for b in pagina if b is not bloque)
            if fijo + alto_bloque(bloque) > COL_H:
                problemas.append(f"{bloque.path.name}: no entra en su página")
            if bloque.w_cm > COL_W_CM + 1e-6 or bloque.h_cm > COL_H_CM + 1e-6:
                problemas.append(f"{bloque.path.name}: excede la caja de texto")
    check("ninguna_imagen_partida_ni_recortada", not problemas, "; ".join(problemas) or "32/32")
    dpis = {m["dpi"] for m in catalogo}
    check("resolucion_minima_200dpi", min(dpis) >= 200, f"dpi: {sorted(dpis)}")

    # 5 bis · el XML respeta el orden de secuencia de OOXML en lo que este generador toca.
    # Word abre igual un archivo con los hijos desordenados, pero avisa de «contenido
    # ilegible» y ofrece repararlo: para un documento que va a Jefatura, eso es un defecto.
    orden = {
        "w:tblPr": ["w:tblStyle", "w:tblpPr", "w:tblOverlap", "w:bidiVisual",
                    "w:tblStyleRowBandSize", "w:tblStyleColBandSize", "w:tblW", "w:jc",
                    "w:tblCellSpacing", "w:tblInd", "w:tblBorders", "w:shd", "w:tblLayout",
                    "w:tblCellMar", "w:tblLook", "w:tblCaption", "w:tblDescription"],
        "w:tblBorders": ["w:top", "w:start", "w:left", "w:bottom", "w:end", "w:right",
                         "w:insideH", "w:insideV"],
    }
    desordenes: list[str] = []
    for padre, secuencia in orden.items():
        for elemento in d.element.body.iter(qn(padre)):
            hijos = [h.tag for h in elemento]
            indices, desconocidos = [], []
            for tag in hijos:
                nombre = next((n for n in secuencia if qn(n) == tag), None)
                (desconocidos if nombre is None else indices).append(nombre or tag)
                if nombre is not None:
                    indices[-1] = secuencia.index(nombre)
            if desconocidos:
                desordenes.append(f"{padre}: hijo inesperado {desconocidos}")
            if indices != sorted(indices):
                desordenes.append(f"{padre}: hijos fuera de secuencia")
            if len(set(indices)) != len(indices):
                desordenes.append(f"{padre}: hijo duplicado")
    check("xml_en_secuencia_ooxml", not desordenes,
          "; ".join(sorted(set(desordenes))) or "tblPr y tblBorders en orden")

    # 5 ter · el esquema. El control de secuencia mira el orden de los hijos y nada más:
    # no vio que `w:zoom` viniera sin su `w:percent`, que OOXML declara requerido. Esto
    # valida cada parte del paquete contra `qa/esquema/`, atributo por atributo.
    errores, n_partes, n_elementos = validar_esquema(docx_path)
    excluidas = ", ".join(f"{p} ({motivo})" for p, motivo in PARTES_EXCLUIDAS.items())
    check("esquema_xsd_ooxml", not errores,
          "; ".join(errores[:4]) if errores else
          f"{n_partes} partes válidas, {n_elementos} elementos distintos cubiertos; "
          f"fuera del alcance, idénticas a la plantilla: {excluidas}")

    # 5 quater · el índice y el folio, que son la manera de moverse por el documento.
    entradas = [b for pagina in paginas for b in pagina if isinstance(b, Idx)]
    donde = {}
    for numero, pagina in enumerate(paginas, start=1):
        for bloque in pagina:
            if isinstance(bloque, P) and bloque.style in ("H1", "H2"):
                donde.setdefault(bloque.runs[0].text, numero)
    desviados = [f"{e.texto}: dice {e.pagina}, está en {donde.get(e.texto)}"
                 for e in entradas if donde.get(e.texto) != e.pagina]
    check("indice_con_folios_reales", entradas and not desviados,
          "; ".join(desviados[:4]) if desviados
          else f"{len(entradas)} entradas, todas en la página que anuncian")

    seccion = d.sections[0]
    pies = seccion._sectPr.findall(qn("w:footerReference"))
    tipos = sorted(p.get(qn("w:type")) for p in pies)
    campos = seccion.footer._element.findall(".//" + qn("w:instrText"))
    portada_limpia = not seccion.first_page_footer._element.findall(".//" + qn("w:instrText"))
    check("folio_al_pie_menos_en_la_portada",
          tipos == ["default", "first"]
          and any("PAGE" in (c.text or "") for c in campos) and portada_limpia,
          f"pies: {tipos}; campo PAGE en el pie corriente: "
          f"{any('PAGE' in (c.text or '') for c in campos)}; portada sin folio: {portada_limpia}")

    # 5 quinquies · las reescrituras de lectura. Ninguna puede quedar sin aplicar —sería una
    # corrección muerta— y ninguna puede mover un número.
    sin_usar = sorted(set(REESCRITURAS) - set(APLICADAS))
    check("reescrituras_aplicadas", not sin_usar,
          f"{len(sin_usar)} sin aplicar: {[s[:40] for s in sin_usar[:3]]}" if sin_usar
          else f"{len(REESCRITURAS)} reescrituras, {sum(APLICADAS.values())} aplicaciones")

    movidas: list[str] = []
    for viejo, nuevo in REESCRITURAS.items():
        antes, despues = NUMERO.findall(viejo), NUMERO.findall(nuevo)
        if antes == despues:
            continue
        faltantes = [n for n in antes if n not in despues]
        agregados = [n for n in despues if n not in antes]
        if all(n in EXCEPCIONES_CIFRAS for n in faltantes) and not agregados:
            continue
        movidas.append(f"{antes} → {despues} en «{viejo[:44]}…»")
    declaradas = "; ".join(f"{n}: {motivo}" for n, motivo in EXCEPCIONES_CIFRAS.items())
    check("diff_de_cifras_contra_conduccion", not movidas,
          "; ".join(movidas[:3]) if movidas
          else f"0 diferencias sin declarar sobre {len(REESCRITURAS)} frases reescritas. "
               f"Declarada: {declaradas}")

    # Y la comprobación de punta a punta: la cifra de cada zona, tal como la publica la
    # edición de conducción, tiene que estar en su ficha y en su fila de la tabla.
    faltan_cifras: list[str] = []
    for rid in B.EXPECTED_REFS:
        ficha = {f["referencia_id"]: f for f in B.CONTENT_CACHE["fichas"]}[rid]
        for numero in NUMERO.findall(C.frase_cifra(ficha)):
            if texto_docx.count(numero) < 2:      # la ficha y la fila de la tabla
                faltan_cifras.append(f"{rid}: {numero}")
    check("cifra_de_cada_zona_en_ficha_y_tabla", not faltan_cifras,
          "; ".join(faltan_cifras) or "22/22")

    # 6 · PDF de prueba
    doc_pdf = fitz.open(pdf_path)
    check("pdf_de_prueba_coincide_en_paginas", doc_pdf.page_count == len(paginas),
          f"{doc_pdf.page_count} páginas")
    doc_pdf.close()

    # 7 · tamaño
    mb = docx_path.stat().st_size / (1 << 20)
    check("tamano_del_docx", mb <= 20.0, f"{mb:.2f} MB")

    campos = ["control", "resultado", "detalle"]
    with destino.open("w", encoding="utf-8", newline="") as fh:
        fh.write(",".join(campos) + "\n")
        for fila in filas:
            fh.write(",".join('"' + str(fila[c]).replace('"', '""') + '"' for c in campos) + "\n")
    return {"filas": filas, "fallas": [f for f in filas if f["resultado"] == "FAIL"]}


# =======================================================================================


def main() -> int:
    parser = argparse.ArgumentParser(
        description="Edición editable (.docx) de la edición de conducción del Atlas DGDGAS.")
    parser.add_argument("--dpi", type=int, default=200,
                        help="Resolución de los mapas exportados (mínimo 200).")
    args = parser.parse_args()
    if args.dpi < 150:
        raise SystemExit("La resolución mínima admitida es 150 dpi.")

    registrar_calibri()
    content = preparar_atlas()

    out = B.OUT
    docx_path = out / "ATLAS_REFERENCIAS_GASTRONOMICAS_CABA_DGDGAS.docx"
    pdf_path = out / "ATLAS_REFERENCIAS_GASTRONOMICAS_CABA_DGDGAS_DOCX.pdf"
    mapas_dir = out / "assets_docx" / f"mapas_{args.dpi}dpi"
    mapas_dir.mkdir(parents=True, exist_ok=True)
    trabajo = out / "assets_docx" / "_trabajo"
    trabajo.mkdir(parents=True, exist_ok=True)
    tmp = trabajo / "_mapa.pdf"

    bloques, catalogo, marcas = construir_documento(content, mapas_dir, args.dpi, tmp)
    bloques, paginas = armar_indice(bloques, marcas)
    emitir_docx(paginas, docx_path)
    usados = emitir_pdf(paginas, pdf_path)
    tmp.unlink(missing_ok=True)
    shutil.rmtree(trabajo, ignore_errors=True)

    reporte = qa(docx_path, pdf_path, paginas, usados, catalogo,
                 out / "qa" / "QA_EDICION_DOCX.csv")

    campos = ["archivo", "clave", "tipo", "ancho_cm", "alto_cm", "px", "dpi", "escala_m", "bytes"]
    with (out / "qa" / "QA_MAPAS_DOCX.csv").open("w", encoding="utf-8", newline="") as fh:
        fh.write(",".join(campos) + "\n")
        for m in catalogo:
            fh.write(",".join(str(m[c]) for c in campos) + "\n")

    resumen = {
        "docx": docx_path.name,
        "docx_bytes": docx_path.stat().st_size,
        "docx_sha256": sha256(docx_path),
        "pdf_prueba": pdf_path.name,
        "pdf_bytes": pdf_path.stat().st_size,
        "pdf_sha256": sha256(pdf_path),
        "paginas": len(paginas),
        "mapas": len(catalogo),
        "dpi": args.dpi,
        "mapas_bytes": sum(m["bytes"] for m in catalogo),
        "pagina_mas_llena_cm": round(max(usados), 2),
        "holgura_minima_cm": round(min(COL_H_CM - u for u in usados), 2),
        "qa_fallas": [f["control"] for f in reporte["fallas"]],
    }
    print(json.dumps(resumen, ensure_ascii=False, indent=2))
    return 1 if reporte["fallas"] else 0


if __name__ == "__main__":
    raise SystemExit(main())
